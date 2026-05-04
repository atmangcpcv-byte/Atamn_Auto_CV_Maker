from docx import Document

doc = Document("C:/Auto_CV_Maker/Resume_template/Resume.docx")

# Meta
doc.paragraphs[0].runs[0].text = "{{ emp_name|upper }}"
for i in range(1, len(doc.paragraphs[0].runs)): doc.paragraphs[0].runs[i].text = ""
doc.paragraphs[1].runs[0].text = "{{ designation }}"
for i in range(1, len(doc.paragraphs[1].runs)): doc.paragraphs[1].runs[i].text = ""
doc.paragraphs[2].runs[0].text = "{{ contact_info }}"
for i in range(1, len(doc.paragraphs[2].runs)): doc.paragraphs[2].runs[i].text = ""
doc.paragraphs[4].runs[0].text = "{{ profile_summary }}"
for i in range(1, len(doc.paragraphs[4].runs)): doc.paragraphs[4].runs[i].text = ""
doc.paragraphs[6].runs[0].text = "{{ skills_string }}"
for i in range(1, len(doc.paragraphs[6].runs)): doc.paragraphs[6].runs[i].text = ""

# Exp
doc.paragraphs[8].insert_paragraph_before("{% for exp in experiences %}")
doc.paragraphs[8].runs[0].text = "{{ exp.title }} | {{ exp.company }}\t{{ exp.duration }}"
for i in range(1, len(doc.paragraphs[8].runs)): doc.paragraphs[8].runs[i].text = ""

doc.paragraphs[9].insert_paragraph_before("{% for b in exp.bullets %}")
doc.paragraphs[9].runs[0].text = "{{ b }}"
doc.paragraphs[10].insert_paragraph_before("{% endfor %}")
doc.paragraphs[10].insert_paragraph_before("{% endfor %}")

# Remove extra exp bullets
doc.paragraphs[10]._element.getparent().remove(doc.paragraphs[10]._element)
doc.paragraphs[11]._element.getparent().remove(doc.paragraphs[11]._element)

# Proj
doc.paragraphs[13].insert_paragraph_before("{% for p in projects %}")
doc.paragraphs[13].runs[0].text = "{{ p.title }}"
for i in range(1, len(doc.paragraphs[13].runs)): doc.paragraphs[13].runs[i].text = ""

doc.paragraphs[14].insert_paragraph_before("{% for b in p.bullets %}")
doc.paragraphs[14].runs[0].text = "{{ b }}"
for i in range(1, len(doc.paragraphs[14].runs)): doc.paragraphs[14].runs[i].text = ""

doc.paragraphs[15].insert_paragraph_before("{% endfor %}")
doc.paragraphs[15].insert_paragraph_before("{% endfor %}")

# Remove extra projects
doc.paragraphs[15]._element.getparent().remove(doc.paragraphs[15]._element)
doc.paragraphs[16]._element.getparent().remove(doc.paragraphs[16]._element)

# Edu
doc.paragraphs[18].insert_paragraph_before("{% for ed in education %}")
doc.paragraphs[18].runs[0].text = "{{ ed.degree }} | {{ ed.institution }}\t{{ ed.graduation_year }}"
for i in range(1, len(doc.paragraphs[18].runs)): doc.paragraphs[18].runs[i].text = ""

doc.paragraphs[19].insert_paragraph_before("{% endfor %}")

# Cert
doc.paragraphs[20].insert_paragraph_before("{% for cert in certifications %}")
doc.paragraphs[20].runs[0].text = "{{ cert }}"
for i in range(1, len(doc.paragraphs[20].runs)): doc.paragraphs[20].runs[i].text = ""

doc.paragraphs[21].insert_paragraph_before("{% endfor %}")
doc.paragraphs[21]._element.getparent().remove(doc.paragraphs[21]._element)

doc.save("C:/Auto_CV_Maker/Resume_template/Resume_Dynamic.docx")
print("Template rebuilt correctly!")
