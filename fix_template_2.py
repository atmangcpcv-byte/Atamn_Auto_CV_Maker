from docx import Document
doc = Document("C:/Auto_CV_Maker/Resume_template/Resume.docx")

def replace_text_in_paragraph(p, search, replace):
    for r in p.runs:
        if search in r.text:
            r.text = r.text.replace(search, replace)

replace_text_in_paragraph(doc.paragraphs[0], "YOUR NAME", "{{ emp_name|upper }}")
replace_text_in_paragraph(doc.paragraphs[1], "Backend Developer", "{{ designation }}")
replace_text_in_paragraph(doc.paragraphs[2], "your.email@example.com", "{{ contact_info }}")
replace_text_in_paragraph(doc.paragraphs[2], "Mobile: +91-XXXXXXXXXX", "")
replace_text_in_paragraph(doc.paragraphs[2], "LinkedIn: linkedin.com/in/yourprofile | City: Your City", "")

# Profile
replace_text_in_paragraph(doc.paragraphs[4], "Motivated backend developer", "{{ profile_summary }}")
for r in doc.paragraphs[4].runs:
    if "Django, REST APIs" in r.text or "Passionate about" in r.text:
         r.text = ""

# Skills
for r in doc.paragraphs[6].runs:
    if "Python | Django | REST APIs" in r.text:
         r.text = "{{ skills_string }}"
    else:
         r.text = ""

# Instead of blindly doing insert_paragraph_before, I will add the tags directly to the strings, then use docxtpl.
# If I append block tags inline, docxtpl handles them.
# Example: 
# PARAGRAPH 7: "EXPERIENCE \n {% for exp in experiences %}" -- docxtpl actually parses newlines in runs as line breaks, not real paragraphs, breaking the loop.
# So I must insert a real paragraph.
p8 = doc.paragraphs[8]
p8.insert_paragraph_before("{% for exp in experiences %}")

for r in p8.runs:
    if "Backend Developer" in r.text:
        r.text = "{{ exp.title }} | {{ exp.company }}"
    elif "2021" in r.text:
        r.text = "\t{{ exp.duration }}"

p9 = doc.paragraphs[9]
p9.insert_paragraph_before("{% for b in exp.bullets %}")
for r in p9.runs:
    if "Developed scalable web" in r.text:
         r.text = "{{ b }}"
    else:
         r.text = ""

p10 = doc.paragraphs[10]
p10.text = "{% endfor %}"
p10.style = "Normal"

p11 = doc.paragraphs[11]
p11.text = "{% endfor %}"
p11.style = "Normal"


p13 = doc.paragraphs[13]
p13.insert_paragraph_before("{% for p in projects %}")
for r in p13.runs:
    if "Inventory Management System" in r.text:
       r.text = "{{ p.title }}"
    else: r.text = ""

p14 = doc.paragraphs[14]
p14.insert_paragraph_before("{% for b in p.bullets %}")
for r in p14.runs:
    if "Built a full-featured inventory" in r.text:
       r.text = "{{ b }}"
    else: r.text = ""

p15 = doc.paragraphs[15]
p15.text = "{% endfor %}"
p15.style = "Normal"

p16 = doc.paragraphs[16]
p16.text = "{% endfor %}"
p16.style = "Normal"


p18 = doc.paragraphs[18]
p18.insert_paragraph_before("{% for ed in education %}")
for r in p18.runs:
    if "Bachelor of Engineering" in r.text:
        r.text = "{{ ed.degree }} | {{ ed.institution }}"
    elif "Graduated Year" in r.text:
        r.text = "\t{{ ed.graduation_year }}"

p19 = doc.paragraphs[19]
p19.insert_paragraph_before("{% endfor %}")

p20 = doc.paragraphs[20]
p20.insert_paragraph_before("{% for cert in certifications %}")
for r in p20.runs:
    if "Certified Django Developer" in r.text:
        r.text = "{{ cert }}"
    else: r.text = ""

p21 = doc.paragraphs[21]
p21.text = "{% endfor %}"
p21.style = "Normal"

doc.save("C:/Auto_CV_Maker/Resume_template/Resume_Dynamic.docx")
print("Done")
