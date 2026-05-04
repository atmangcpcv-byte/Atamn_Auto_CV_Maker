from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

# Load original to keep styles/margins
doc = Document("C:/Auto_CV_Maker/Resume_template/Resume.docx")

# Clear existing content safely
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)

def add_heading(text):
    p = doc.add_paragraph()
    p.add_run().add_break()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Add a bottom border or underline
    r.underline = True

# Header
p_name = doc.add_paragraph()
r = p_name.add_run("{{ emp_name|upper }}")
r.bold = True
r.font.size = Pt(16)
p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p_desig = doc.add_paragraph()
p_desig.add_run("{{ designation }}").font.size = Pt(12)
p_desig.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p_contact = doc.add_paragraph()
p_contact.add_run("{{ contact_info }}").font.size = Pt(10)
p_contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Summary
add_heading("PROFILE SUMMARY")
doc.add_paragraph("{{ profile_summary }}")

# Skills
add_heading("SKILLS")
doc.add_paragraph("{{ skills_string }}")

# Experience
add_heading("EXPERIENCE")
p_exp_loop = doc.add_paragraph()
p_exp_loop.add_run("{% for exp in experiences %}")

p_exp_title = doc.add_paragraph()
r_title = p_exp_title.add_run("{{ exp.title }} | {{ exp.company }}")
r_title.bold = True
r_title.add_tab()
r_title.add_tab()
r_date = p_exp_title.add_run("{{ exp.duration }}")
r_date.bold = True

p_exp_bullet = doc.add_paragraph(style='List Bullet')
p_exp_bullet.add_run("{% for b in exp.bullets %}")

p_bullet = doc.add_paragraph("{{ b }}", style='List Bullet')

p_end_bullet = doc.add_paragraph()
p_end_bullet.add_run("{% endfor %}")

p_end_exp = doc.add_paragraph()
p_end_exp.add_run("{% endfor %}")

# Projects
add_heading("PROJECTS")
doc.add_paragraph("{% for p in projects %}")
p_proj_title = doc.add_paragraph()
p_proj_title.add_run("{{ p.title }}").bold = True
doc.add_paragraph("{% for b in p.bullets %}", style='List Bullet')
doc.add_paragraph("{{ b }}", style='List Bullet')
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endfor %}")

# Education
add_heading("EDUCATION")
doc.add_paragraph("{% for ed in education %}")
p_ed = doc.add_paragraph()
p_ed.add_run("{{ ed.degree }} | {{ ed.institution }}").bold = True
p_ed.add_run("\t{{ ed.graduation_year }}")
doc.add_paragraph("{% endfor %}")

# Certifications
add_heading("CERTIFICATIONS")
doc.add_paragraph("{% for cert in certifications %}")
doc.add_paragraph("{{ cert }}", style='List Bullet')
doc.add_paragraph("{% endfor %}")

doc.save("C:/Auto_CV_Maker/Resume_template/Resume_Dynamic.docx")
print("Dynamic template generated successfully!")
